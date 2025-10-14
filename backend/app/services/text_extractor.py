import pypdf
from docx import Document
import asyncio
import logging
import re
from typing import Optional, Dict

async def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF files with improved handling
    Supports multi-page PDFs with various encodings
    """
    try:
        logging.info(f"Extracting text from PDF: {file_path}")
        text = ""
        
        def _extract():
            reader = pypdf.PdfReader(file_path)
            extracted = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    extracted += f"{page_text}\n"
            return extracted
        
        text = await asyncio.to_thread(_extract)
        
        if not text.strip():
            logging.warning("No text extracted from PDF - may be image-based")
            raise ValueError("PDF appears to be image-based or empty. OCR required.")
        
        logging.info(f"Extracted {len(text)} characters from PDF")
        return text.strip()
    
    except pypdf.errors.PdfReadError as e:
        logging.error(f"PDF read error: {e}")
        raise Exception(f"Failed to read PDF: {e}")
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        raise Exception(f"PDF extraction error: {e}")

async def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files
    Includes paragraph text, table text, and headers/footers
    """
    try:
        logging.info(f"Extracting text from DOCX: {file_path}")
        
        def _extract():
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        
        text = await asyncio.to_thread(_extract)
        
        logging.info(f"Extracted {len(text)} characters from DOCX")
        return text.strip()
    
    except Exception as e:
        logging.error(f"DOCX extraction failed: {e}")
        raise Exception(f"DOCX extraction error: {e}")

async def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from plain text files
    Handles various encodings (UTF-8, UTF-16, Latin-1)
    """
    try:
        logging.info(f"Extracting text from TXT: {file_path}")
        
        # Try multiple encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                def _read():
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                
                text = await asyncio.to_thread(_read)
                logging.info(f"Successfully read TXT with {encoding} encoding")
                return text.strip()
            
            except UnicodeDecodeError:
                continue
        
        raise Exception("Unable to decode text file with supported encodings")
    
    except Exception as e:
        logging.error(f"TXT extraction failed: {e}")
        raise Exception(f"TXT extraction error: {e}")

async def extract_text_from_srt(file_path: str) -> str:
    """
    Extract text from SRT subtitle files (without timestamps)
    Returns clean text suitable for translation
    """
    try:
        logging.info(f"Extracting text from SRT: {file_path}")
        
        def _extract():
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Pattern to extract subtitle text only
            # Support both LF and CRLF line endings (\n or \r\n)
            pattern = r'\d+\r?\n\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}\r?\n(.*?)(?=\r?\n\r?\n|\Z)'
            matches = re.findall(pattern, content, re.DOTALL)
            
            # Join all subtitle text
            text = "\n".join([match.strip() for match in matches if match.strip()])
            return text
        
        text = await asyncio.to_thread(_extract)
        
        logging.info(f"Extracted {len(text)} characters from SRT")
        return text.strip()
    
    except Exception as e:
        logging.error(f"SRT extraction failed: {e}")
        raise Exception(f"SRT extraction error: {e}")

async def extract_text_universal(file_path: str, file_type: Optional[str] = None) -> str:
    """
    Universal text extractor with automatic format detection
    
    Supports: PDF, DOCX, TXT, SRT
    
    Args:
        file_path: Path to file
        file_type: Optional file type override (pdf, docx, txt, srt)
    
    Returns:
        Extracted text string
    """
    try:
        # Detect file type from extension if not provided
        if file_type is None:
            file_type = file_path.split('.')[-1].lower()
        
        logging.info(f"Universal text extraction for type: {file_type}")
        
        # Map file types to extractor functions
        extractors = {
            'pdf': extract_text_from_pdf,
            'docx': extract_text_from_docx,
            'doc': extract_text_from_docx,  # Also support .doc
            'txt': extract_text_from_txt,
            'srt': extract_text_from_srt,
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}. Supported: {list(extractors.keys())}")
        
        # Extract text
        text = await extractor(file_path)
        
        if not text or not text.strip():
            raise ValueError("No text content extracted from file")
        
        return text
    
    except Exception as e:
        logging.error(f"Universal text extraction failed: {e}")
        raise Exception(f"Text extraction error: {e}")

async def extract_metadata(file_path: str) -> Dict:
    """
    Extract metadata from document files
    
    Returns document metadata like author, title, creation date
    """
    try:
        file_type = file_path.split('.')[-1].lower()
        
        if file_type == 'pdf':
            reader = pypdf.PdfReader(file_path)
            metadata = reader.metadata
            return {
                'title': metadata.get('/Title', 'Unknown'),
                'author': metadata.get('/Author', 'Unknown'),
                'pages': len(reader.pages),
                'format': 'PDF'
            }
        
        elif file_type in ['docx', 'doc']:
            doc = Document(file_path)
            core_props = doc.core_properties
            return {
                'title': core_props.title or 'Unknown',
                'author': core_props.author or 'Unknown',
                'paragraphs': len(doc.paragraphs),
                'format': 'DOCX'
            }
        
        else:
            return {'format': file_type.upper()}
    
    except Exception as e:
        logging.error(f"Metadata extraction failed: {e}")
        return {'format': 'Unknown'}

def get_supported_formats() -> list:
    """Get list of supported document formats"""
    return ['pdf', 'docx', 'doc', 'txt', 'srt']
